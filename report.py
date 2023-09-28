###Known bugs
#1. Detecting the line that has the RP and DMI in comments
#2. In composite, show ELMOD values for both concrete and asphalt. Currently showing only asphalt


from pydoc import Doc
from venv import create
from webbrowser import get
import docx
from docx import Document
from docx.shared import Inches
from datetime import date, datetime
import os
import pickle
import matplotlib.pyplot as plt
from matplotlib.ticker import AutoMinorLocator, FormatStrFormatter
import numpy as np
import excel, comments, report_page4
from docx.enum.text import WD_ALIGN_PARAGRAPH
from query_db import *

# if os.path.exists("TestDoc.docx"):
#   os.remove("TestDoc.docx")
# else:
#   print("The file does not exist") 

def elab_dir(dir):
    if(dir == "EB"):
        return "East Bound"
    elif(dir == "NB"):
        return "North Bound"
    elif(dir == "SB"):
        return "South Bound"
    elif(dir == "WB"):
        return "West Bound"

def create_doc():
    document = Document()
    section = document.sections[-1]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "Left text\tCenter Text\tRight Text"
    paragraph.style = document.styles["Header"]
    document.save("TestDoc.docx")

def set_header_footer(mde_path, ll, treatment):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = docx.shared.Pt(12)

    section = document.sections[-1]
    header = section.header
    paragraph = header.paragraphs[0]
    # paragraph.text = "Left text\tCenter Text\tRight Text"
    today = date.today()
    d1 = today.strftime("%m/%d/%Y")
    header_txt = "R&D ID No: "+ll["req no"]+"\t\tPage Number here\n"+d1
    paragraph.text = header_txt
    paragraph.style = document.styles["Header"]

    footer = section.footer
    paragraph = footer.paragraphs[0]
    # paragraph.text = "Left text\tCenter Text\tRight Text"
    # paragraph.text = ll["roadname"]+" from "+(ll["rp from"].split('.')[0])+"+"+(ll["rp from"].split('.')[1])+" to "+(ll["rp to"].split('.')[0])+"+"+(ll["rp to"].split('.')[1])+"\t\t"+treatment
    if(ll["pavtype"] == "asphalt"):
        treatment = "Overlay Design"
    elif(ll["pavtype"] == "concrete"):
        treatment = "Underseal"
    elif(ll["pavtype"] == "composite"):
        treatment = "Underseal and Overlay Design"
    paragraph.text = get_pathstring(mde_path) + "\t\t" + treatment
    paragraph.style = document.styles["footer"]
    return document

    # document.save("TestDoc.docx")

def get_thickness(mde, pavtype):
    arr = np.array(mde["mod_est"])
    arr = arr[:, [3, 4, 5, 6]]
    arr = arr.astype(float)
    arr = np.transpose(arr)
    thickness_str = ""
    if (np.all(arr[1:4, :] == 0)):
        layers = 2
        if (pavtype == "asphalt"):
            thickness_str = str(arr[0][0])+" in. Asphalt\n   --         Subgrade"
            # limits = [mod_limits["asphalt"], mod_limits["subgrade"]]
        elif (pavtype == "concrete"):
            # limits = [mod_limits["concrete"], mod_limits["subgrade"]]
            thickness_str = str(arr[0][0])+" in. Concrete\n  --         Subgrade"
    elif (np.all(arr[2:4, :] == 0)):
        layers = 3
        if (pavtype == "asphalt"):
            thickness_str = str(arr[0][0])+" in. Asphalt\n"+str(arr[1][0])+" in. Subbase\n  --         Subgrade"
            # limits = [mod_limits["asphalt"], mod_limits["middle"], mod_limits["subgrade"]]
        elif (pavtype == "concrete"):
            thickness_str = str(arr[0][0])+" in. Concrete\n"+str(arr[1][0])+" in. Subbase\n  --         Subgrade"
        elif (pavtype == "composite"):
            thickness_str = str(arr[0][0])+" in. Asphalt\n"+str(arr[1][0])+" in. Concrete\n  --         Subgrade"
    elif (np.all(arr[3:4] == 0)):
        layers = 4
        thickness_str = str(arr[0][0])+" in. Asphalt\n"+str(arr[1][0])+" in. Concrete\n"+str(arr[1][0])+" in. Subbase"+"\n  --         Subgrade"
    return thickness_str

def align_middle_cells(cells):
    for cell in cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return cells

def cover_page(ll_no, year, mde_path, ll, mde, document,con):
    heading = ll["roadname"] + " Pavement Falling Weight Deflectometer (FWD) Testing Results"
    # document.add_heading(heading, 1)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence = p.add_run(heading)
    sentence.font.name = 'Arial'
    sentence.font.size = docx.shared.Pt(16)
    sentence.font.underline = True
    sentence.bold = True
    sentence.font.color.rgb = docx.shared.RGBColor(128, 0, 128)

    p = document.add_paragraph()

    if(ll["pavtype"] == "asphalt"):
        sentence = p.add_run("Overlay Design")
    elif(ll["pavtype"] == "concrete"):
        sentence = p.add_run("Underseal")
    elif(ll["pavtype"] == "composite"):
        sentence = p.add_run("Underseal and Overlay Design")
    
    # sentence = p.add_run(heading)
    sentence.font.name = 'Arial'
    sentence.font.size = docx.shared.Pt(14)
    sentence.font.underline = True
    sentence.bold = True
    sentence.font.color.rgb = docx.shared.RGBColor(26, 36, 238)
    # document.add_paragraph("\n")

    route_table = document.add_table(rows = 4, cols = 2)
    route_table.autofit = False
    route_table.autofit = False
    route_table.style = 'Table Grid'
    cells = route_table.rows[0].cells

    # cells = align_middle_cells(cells)
    # cells[0].text = "Route: "
    # cells[1].text = ll["roadname"]+" from "+(ll["rp from"].split('.')[0])+"+"+(ll["rp from"].split('.')[1])+" to "+(ll["rp to"].split('.')[0])+"+"+(ll["rp to"].split('.')[1])
    cell_width = 2
    cell_width2 = 4
    run = cells[0].paragraphs[0].add_run("Route: ")
    cells[0].width = Inches(cell_width)
    run.bold = True
    cells[1].text = get_pathstring(mde_path)
    cells[1].width = Inches(cell_width2)
    #ll["roadname"]+" from "+ll["rp from"]+" to "+ll["rp to"]

    cells = route_table.rows[1].cells
    # cells = align_middle_cells(cells)
    run = cells[0].paragraphs[0].add_run("District: ")
    cells[0].width = Inches(cell_width)
    run.bold = True
    cells[1].text = ll["district"]
    cells[1].width = Inches(cell_width2)

    cells = route_table.rows[2].cells
    # cells = align_middle_cells(cells)
    # cells[0].text = "Contract Number: "
    run = cells[0].paragraphs[0].add_run("Contract Number: ")
    cells[0].width = Inches(cell_width)
    run.bold = True
    cells[1].text = ll["des no"]
    cells[1].width = Inches(cell_width2)

    cells = route_table.rows[3].cells
    # cells = align_middle_cells(cells)
    # cells[0].text = "R&D Testing Request ID Number: "
    run = cells[0].paragraphs[0].add_run("R&D Testing Request\nID Number: ")
    cells[0].width = Inches(cell_width)
    run.bold = True
    cells[1].text = ll["req no"]
    cells[1].width = Inches(cell_width2)

    document.add_paragraph(" ")

    testing_table = document.add_table(rows = 1, cols = 2)
    testing_table.style = 'Table Grid'
    cells = testing_table.rows[0].cells
    run = cells[0].paragraphs[0].add_run("Testing conducted on: ")
    cells[0].width = Inches(cell_width)
    run.bold = True
    testing_date_time = mde['deflections'][0][2] ## [0] refers to first row and [2] refers to the 'TheTime' column
    # print('tesing date:',testing_date)
    # print('type pf testdate:',type(testing_date))
    # Refine the time format
    testing_date = testing_date_time.split(' ')[0]
    test_date_splited = testing_date.split('-')
    test_year = test_date_splited[0]
    test_month = test_date_splited[1]
    test_day = test_date_splited[2]
    cells[1].text = "{}/{}/{}".format(test_month,test_day,test_year)
    cells[1].width = Inches(cell_width2)

    document.add_paragraph(" ")

    analysis_table = document.add_table(rows = 2, cols = 3)
    analysis_table.style = 'Table Grid'
    # analysis_table.autofit = True
    cells = analysis_table.rows[0].cells
    run = cells[0].paragraphs[0].add_run("Analysis Type requested: ")
    run.bold = True
    cells[1].merge(cells[2])
    if(ll["pavtype"] == "asphalt"):
        run = cells[1].paragraphs[0].add_run("Overlay Design")
    elif(ll["pavtype"] == "concrete"):
        run = cells[1].paragraphs[0].add_run("Underseal")
        # cells[1].text = "Underseal"
    elif(ll["pavtype"] == "composite"):
        run = cells[1].paragraphs[0].add_run("Underseal and Overlay Design")
        # cells[1].text = "Underseal and Overlay Design"
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cells = analysis_table.rows[1].cells
    run = cells[0].paragraphs[0].add_run("Pavement Profile used in Analysis: ")
    run.bold = True
    # [:2] is to extract the first 2 letter from the dir string ('WBDL' for west bound driving lane)
    # in 2022 longlist the direction includes the lane name
    # So 'WB' will be changed to one of ['WBDL','WBPL','WBSH']
    run = cells[1].paragraphs[0].add_run(elab_dir(ll["dir"][:2])) 
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cells[2].paragraphs[0].add_run(get_thickness(mde, ll["pavtype"]))
    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return document

def underseal_page(rp_str, mde_path, ll, calc_data, document):
    # p = document.add_paragraph()
    from_rp, to_rp = report_page4.get_from_rp_to_rp_str(mde_path)
    head = elab_dir(ll["dir"][:2]) + " Lane from RP-" + from_rp[0]+ "+" + from_rp[1] + " to RP-" + to_rp[0] + "+" +to_rp[1] + "\n"
    # p.add_run(heading).bold = True
    heading = document.add_heading()
    sentence = heading.add_run(head)
    # heading = document.add_heading(head, 3)
    sentence.font.name = 'Arial'
    sentence.font.size = docx.shared.Pt(13)
    sentence.font.underline = True
    sentence.bold = True
    sentence.font.color.rgb = docx.shared.RGBColor(128, 0, 128)
    line = "Underseal requirement profile, " + rp_str
    p = document.add_paragraph()
    p.add_run(line)
    
    table = document.add_table(rows = 1, cols = 5)
    table.style = 'Table Grid'
    cells = table.rows[0].cells
    cells = align_middle_cells(cells)
    cells[0].paragraphs[0].add_run("From DMI (ft)")
    cells[1].paragraphs[0].add_run("To DMI (ft)")
    cells[2].paragraphs[0].add_run("Length (ft)")
    cells[3].paragraphs[0].add_run("HMA MODULUS (ksi)")
    cells[4].paragraphs[0].add_run("PCC MODULUS (ksi)")
    underseal_total = 0
    for i in range(1, (len(calc_data["invalid_table"])+1)):
        # print("Came here")
        # print(calc_data["invalid_table"][i-1])
        cells = table.add_row().cells
        cells = align_middle_cells(cells)
        for j in range(5):
            # print(j)
            # print(calc_data["invalid_table"][i-1][j])
            # cells[j].text = "a"
            cells[j].paragraphs[0].add_run(str(calc_data["invalid_table"][i-1][j]))
        underseal_total = underseal_total + calc_data["invalid_table"][i-1][2]
    
    document.add_paragraph(" ")
    table = document.add_table(rows = 5, cols = 3)
    table.style = "Table Grid"
    cells = table.rows[0].cells
    run = cells[0].paragraphs[0].add_run("Total length of pavement tested: ")
    run.bold = True
    # length = round((abs(float(ll["rp from"]) - float(ll["rp to"])))*5280)
    length = calc_dist(mde_path)
    run = cells[1].paragraphs[0].add_run(str(length))
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[2].text = "feet"

    cells = table.rows[1].cells
    run = cells[0].paragraphs[0].add_run("Total Length of Underseal: ")
    run.bold = True
    run = cells[1].paragraphs[0].add_run(str(underseal_total))
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[2].text = "feet"

    cells = table.rows[2].cells
    run = cells[0].paragraphs[0].add_run("Amount of Underseal required: ")
    run.bold = True
    run = cells[1].paragraphs[0].add_run(str(round(0.3048*8.68799973*underseal_total/1000, 3)))
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[2].text = "megagrams"

    cells = table.rows[3].cells
    run = cells[0].paragraphs[0].add_run("% of pavement to be undersealed: ")
    run.bold = True
    run = cells[1].paragraphs[0].add_run(str(round((underseal_total/length) * 100, 2)))
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[2].text = "%"

    cells = table.rows[4].cells
    run = cells[0].paragraphs[0].add_run("% of pavement not to be undersealed: ")
    run.bold = True
    run = cells[1].paragraphs[0].add_run(str(round(((length-underseal_total)/length)*100, 2)))
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[2].text = "%"

    document.add_paragraph(" ")

    table = document.add_table(rows = 3, cols = 1)
    table.style = "Table Grid"
    cells = table.rows[0].cells
    run = cells[0].paragraphs[0].add_run("* This material estimates are based on 16 holes per every 100 meter of underseal and approximately 12 gallon of underseal material per hole.")
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(12)
    # run.font.underline = True
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(26, 36, 238)


    cells = table.rows[1].cells
    run = cells[0].paragraphs[0].add_run("** Please note that underseal testing is performed in the driving lane, if you have two lane road the amount of underseal materials need to be multiplied by 2.")
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(12)
    # run.font.underline = True
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(26, 36, 238)

    cells = table.rows[2].cells
    run = cells[0].paragraphs[0].add_run("*** Please check with the elastic modulus of concrete above.  If the modulus is less than 800 ksi, it means that the concrete has cracks around the joints, underseal may not be a good option.  ")
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(12)
    # run.font.underline = True
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(237, 125, 49)

    document.add_page_break()

    return document

def chart(rp_str, stats_data, calc_data, mde, document):
    # print('[Chart] mde deflections chainage:{}'.format(mde['deflections'][:,1]))
    d0_crit = stats_data["d0crit"]
    subgd_crit = stats_data["subgd_calc"]
    # print(calc_data["sensordata"])
    sensordata = np.array(calc_data["sensordata"])
    # print(sensordata)
    # print(sensordata.shape)
    d0 = sensordata[:, 0]
    # print(d0)
    d8 = sensordata[:, 8]
    # print(d8)
    d0_crit_arr = len(d0)*[d0_crit] 
    subgd_crit_arr = len(d8)*[subgd_crit]

    # print(len(d0))
    # print(len(d8))
    # print(len(d0_crit_arr))
    # print(len(subgd_crit_arr))
    

    chainages = []
    for i in range(len(mde["deflections"])):
        if(int(mde["deflections"][i][4]) == 2):
            chainages.append(float(mde["deflections"][i][1]))
            chainages[-1] = chainages[-1]*3.28084 
    # print(len(chainages))

    # Calcualte the max value of chart
    max_value = 0
    for d_arr in [d0,d8,d0_crit_arr,subgd_crit_arr]:
        d_arr_max = np.amax(d_arr)
        if d_arr_max > max_value:
            max_value = d_arr_max

    ### some attributes that can be tunned in plots
    # This factor aims at adjusting the max value along y_axis
    enlarge_factor = 1.7
    # These variables can adjust the font size of different part in the chart
    legend_font_size = 9
    super_title_font_size = 20
    axis_label_fontsize = 16
    x_tick_font_size = 4
    target_num_ticks = 20
    x_tick_dilate_rate = len(chainages)//target_num_ticks + 1 # +1 to prevent dilate rate of 0
    x_label_distance = 20 # distance in points between x_ticks and label of x axis

    # Start plotting the chart
    fig,ax = plt.subplots(figsize= (10,7))
    ax.plot(chainages, d0_crit_arr, label="Surface Deflection Criteria", linewidth=2.0, linestyle='--', color='red')
    ax.plot(chainages, subgd_crit_arr, label="Subgrade Deflection Criteria", linewidth=2.0, linestyle='--', color='magenta')
    ax.plot(chainages, d0, marker='d', markersize=3, label="Surface Deflection", linewidth=1.0, color='blue')
    ax.plot(chainages, d8, marker='.', markersize=3, label="Subgrade Deflection", linewidth=1.0, color='green')
    fig.suptitle("Surface and Subgrade Deflection",y=0.98, fontsize = super_title_font_size)
    ax.set_xlabel("FWD Station, DMI (feet)", fontsize=axis_label_fontsize, labelpad=x_label_distance)
    ax.set_ylabel("Deflection (mils)", fontsize=axis_label_fontsize)
    ax.set_ylim((0.0, enlarge_factor*max_value))
    if x_tick_dilate_rate > 1:
        ax.set_xticks(np.arange(min(chainages), 
                                max(chainages)+1, 
                                (max(chainages)-min(chainages))/(len(chainages)//x_tick_dilate_rate)))
    ax.set_xticklabels(ax.get_xticks(), rotation = 90)
    ax.xaxis.set_major_formatter(FormatStrFormatter('%d'))
    ax.xaxis.set_minor_locator(AutoMinorLocator(5))
    ax.grid(axis="y")
    ax.legend(bbox_to_anchor=(0,1.02,1,0.2), loc="lower left",
              mode="expand", borderaxespad=0, ncol=4, fontsize=legend_font_size)
    fig.savefig("fig.png", bbox_inches = 'tight', dpi = 600)
    plt.close()

    p = document.add_paragraph()
    heading = "Deflection profile, " + rp_str
    p.add_run(heading)#.bold = True
    p = document.add_paragraph("\n")
    p = document.add_picture("fig.png", width = Inches(6))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_page_break()
    return document
    # plt.show()

def apply_font(table, size, rows):
    for i in range(rows):
        cells = table.rows[i].cells
        for cell in cells:
            runs = cell.paragraphs[0].runs
            for run in runs:
                run.font.size = docx.shared.Pt(10)
    
    return table

def esals_page(rp_str, calc_data, document):
    esals_sheet = calc_data["esals_sheet"]

    p = document.add_paragraph()
    heading = "Estimated Remaining Life, " + rp_str + "\n"
    p.add_run(heading)#.bold = True
    
    # p = document.add_paragraph()
    run = p.add_run("Pavement Remaining Life based on Structural Number")
    run.bold = True
    run.font.size = docx.shared.Pt(10)
    # p.add_run("Harmonic Mean").bold = True
    
    table = document.add_table(rows = 4, cols = 3)
    table.style = "Table Grid"
    cells = table.rows[0].cells
    cells[0].merge(cells[2])
    run = cells[0].paragraphs[0].add_run("Harmonic Mean")
    run.bold = True
    
    cells = table.rows[1].cells
    run = cells[0].paragraphs[0].add_run("Structural Number")
    run.bold = True
    run = cells[1].paragraphs[0].add_run(str(round(esals_sheet["hmeans"]["sn"], 2)))
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    cells = table.rows[2].cells
    run = cells[0].paragraphs[0].add_run("Remaining ESALs (50% reliabaility)")
    run.font.color.rgb = docx.shared.RGBColor(83, 129, 53)
    run.bold = True
    run = cells[1].paragraphs[0].add_run(str(round(esals_sheet["hmeans"]["remesals"])))
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.font.color.rgb = docx.shared.RGBColor(83, 129, 53)
    run = cells[2].paragraphs[0].add_run("ESALs")
    run.font.color.rgb = docx.shared.RGBColor(83, 129, 53)

    cells = table.rows[3].cells
    run = cells[0].paragraphs[0].add_run("Modulus of Resilience")
    run.bold = True
    run = cells[1].paragraphs[0].add_run(str(round(esals_sheet["hmeans"]["mod_res"])))
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[2].text = "psi"

    table = apply_font(table, 10, 4)
    p = document.add_paragraph("\n")
    run = p.add_run("Pavement Remaining Life based on Roughness")
    run.bold = True
    run.font.size = docx.shared.Pt(10)

    table = document.add_table(rows = 3, cols = 3)
    table.style = "Table Grid"
    cells = table.rows[0].cells
    run = cells[0].paragraphs[0].add_run("Present IRI (English unit)")
    run.bold = True
    run = cells[1].paragraphs[0].add_run(str(esals_sheet["iri"]))
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[2].text = "inch/mile"

    cells = table.rows[1].cells
    run = cells[0].paragraphs[0].add_run("Present IRI (Metric unit)")
    run.bold = True
    run = cells[1].paragraphs[0].add_run(str(esals_sheet["iri_metric"]))
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[2].text = "meters/km"

    cells = table.rows[2].cells
    run = cells[0].paragraphs[0].add_run("Estimated Present PSI")
    run.bold = True
    run = cells[1].paragraphs[0].add_run(str(esals_sheet["estimated psi"]))
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[2].text = ""
    table = apply_font(table, 10, 3)

    p = document.add_paragraph("")
    table = document.add_table(rows = 5, cols = 4)
    table.style = "Table Grid"
    cells = table.rows[0].cells
    run = cells[0].paragraphs[0].add_run("Baseline ESALs to PSI=2.5")
    run.bold = True
    cells[1].text = str(round(esals_sheet["baseline esals"]))
    
    cells = table.rows[1].cells
    run = cells[0].paragraphs[0].add_run("So coefficient")
    run.bold = True
    cells[1].text = str(0.35)
    
    cell_width = Inches(4)
    cells = table.rows[2].cells
    run = cells[0].paragraphs[0].add_run("Remaining ESALs, 95% reliability")
    cells[0].width = cell_width
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    run = cells[1].paragraphs[0].add_run(str(round(esals_sheet["esal 95"])))
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    run = cells[2].paragraphs[0].add_run("ESALs")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    run = cells[3].paragraphs[0].add_run("Interstates")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)

    cells = table.rows[3].cells
    run = cells[0].paragraphs[0].add_run("Remaining ESALs, 90% reliability")
    cells[0].width = cell_width
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    run = cells[1].paragraphs[0].add_run(str(round(esals_sheet["esal 90"])))
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    run = cells[2].paragraphs[0].add_run("ESALs")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    run = cells[3].paragraphs[0].add_run("US Highways")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)

    cells = table.rows[4].cells
    run = cells[0].paragraphs[0].add_run("Remaining ESALs, 80% reliability")
    cells[0].width = cell_width
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    run = cells[1].paragraphs[0].add_run(str(round(esals_sheet["esal 80"])))
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    run = cells[2].paragraphs[0].add_run("ESALs")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    run = cells[3].paragraphs[0].add_run("State Routes")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)

    table = apply_font(table, 10, 5)

    document.add_paragraph("")
    document = esals_default_table(document)
    document.add_page_break()
    return document

def get_pathstring(mde_path):
    tmp = os.path.splitext(os.path.basename(mde_path))[0]
    arr = tmp.split()
    arr[1] = "from"
    tmp = ' '.join(arr)
    return tmp

def calc_dist(mde_path):
    """
    (TODO)
    Maybe we can use regex to extract this. Modify it later.
    """
    path = get_pathstring(mde_path)
    tmp = path.split()
    d1 = tmp[2]
    d2 = tmp[4]
    rp1 = 0
    dmi1 = 0
    neg1 = None
    till = 0
    for i in range(len(d1[3:])):
        if(d1[3+i] == '+'):
            till = 3+i
            break
        rp1 = rp1*10+(int(d1[3+i]))
    if(d1[till+1] == '-'):
        neg1 = 1
        till = till + 1
    # print("till: ", till)
    dmi1 = float(d1[till+1:len(d1)])
    # # print('dmi1_wen',dmi1)
    # for i in range(till+1, len(d1)):
    #     dmi1 = dmi1*10+(int(d1[i]))
    # print('dmi1',dmi1)
#########################################
    rp2 = 0
    dmi2 = 0
    neg2 = None
    till = 0
    for i in range(len(d2[3:])):
        if(d2[3+i] == '+'):
            till = 3+i
            break
        rp2 = rp2*10+(int(d2[3+i]))
    if(d2[till+1] == '-'):
        neg2 = 1
        till = till + 1
    # print("till: ", till)
    dmi2 = float(d2[till+1:len(d2)])
    # # print('dmi1_wen',dmi1)
    # for i in range(till+1, len(d1)):
    #     dmi1 = dmi1*10+(int(d1[i]))
    # print('dmi2',dmi2)

    dist1 = 0
    dist2 = 0
    if(neg1 == 1):
        dist1 = int(rp1) - (int(dmi1)/100.0)
    else:
        dist1 = int(rp1) + (int(dmi1)/100.0)
    
    if(neg2 == 1):
        dist2 = int(rp2) - (int(dmi2)/100.0)
    else:
        dist2 = int(rp2) + (int(dmi2)/100.0)

    # print("dist1: ", dist1, "\ndist2: ", dist2)
    return round(abs((dist2*5280)-(dist1*5280)))
    

def esals_default_table(document):
    table = document.add_table(rows = 17, cols = 4)
    table.style = "Table Grid"
    cells = table.rows[0].cells
    cells[1].merge(cells[3])
    run = cells[1].paragraphs[0].add_run("Typical Accumulated ESALs")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)

    cells = table.rows[1].cells
    run = cells[0].paragraphs[0].add_run("AADT")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[1].paragraphs[0].add_run("1 year")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[2].paragraphs[0].add_run("5 years")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[3].paragraphs[0].add_run("10 years")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)

    cells = table.rows[2].cells
    run = cells[0].paragraphs[0].add_run("1,000")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[1].paragraphs[0].add_run("44,165")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[2].paragraphs[0].add_run("239,212")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[3].paragraphs[0].add_run("530,250")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)

    cells = table.rows[3].cells
    run = cells[0].paragraphs[0].add_run("2,000")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[1].paragraphs[0].add_run("88,330")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[2].paragraphs[0].add_run("478,424")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[3].paragraphs[0].add_run("1,060,499")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)

    cells = table.rows[4].cells
    run = cells[0].paragraphs[0].add_run("5,000")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[1].paragraphs[0].add_run("220,825")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[2].paragraphs[0].add_run("1,196,059")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[3].paragraphs[0].add_run("2,651,249")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)

    cells = table.rows[5].cells
    run = cells[0].paragraphs[0].add_run("7,500")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[1].paragraphs[0].add_run("331,238")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[2].paragraphs[0].add_run("1,794,089")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[3].paragraphs[0].add_run("3,976,873")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)

    cells = table.rows[6].cells
    run = cells[0].paragraphs[0].add_run("10,000")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[1].paragraphs[0].add_run("441,650")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[2].paragraphs[0].add_run("2,392,119")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[3].paragraphs[0].add_run("5,302,497")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)

    cells = table.rows[7].cells
    run = cells[0].paragraphs[0].add_run("20,000")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[1].paragraphs[0].add_run("883,300")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[2].paragraphs[0].add_run("4,784,238")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[3].paragraphs[0].add_run("10,604,994")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)

    cells = table.rows[8].cells
    run = cells[0].paragraphs[0].add_run("50,000")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[1].paragraphs[0].add_run("2,208,250")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[2].paragraphs[0].add_run("11,960,594")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[3].paragraphs[0].add_run("26,512,486")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)


    cells = table.rows[9].cells
    run = cells[0].paragraphs[0].add_run("75,000")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[1].paragraphs[0].add_run("3,312,375")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[2].paragraphs[0].add_run("17,940,891")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[3].paragraphs[0].add_run("39,768,729")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)

    cells = table.rows[10].cells
    run = cells[0].paragraphs[0].add_run("100,000")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[1].paragraphs[0].add_run("4,416,500")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[2].paragraphs[0].add_run("23,921,189")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[3].paragraphs[0].add_run("53,024,972")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)

    cells = table.rows[11].cells
    run = cells[0].paragraphs[0].add_run("125,000")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[1].paragraphs[0].add_run("5,520,625")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[2].paragraphs[0].add_run("29,901,486")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[3].paragraphs[0].add_run("66,281,215")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)

    cells = table.rows[12].cells
    run = cells[0].paragraphs[0].add_run("150,000")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[1].paragraphs[0].add_run("6,624,750")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[2].paragraphs[0].add_run("35,881,783")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)
    run = cells[3].paragraphs[0].add_run("79,537,458")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)



    cells = table.rows[13].cells
    cells[0].merge(cells[3])
    run = cells[0].paragraphs[0].add_run("Based on: ")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)

    cells = table.rows[14].cells
    cells[0].merge(cells[3])
    run = cells[0].paragraphs[0].add_run("Single unit trucks = 5%")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)


    cells = table.rows[15].cells
    cells[0].merge(cells[3])
    run = cells[0].paragraphs[0].add_run("Semi Trailers = 7%")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)

    cells = table.rows[16].cells
    cells[0].merge(cells[3])
    run = cells[0].paragraphs[0].add_run("Annual Traffic Growth = 4%")
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)

    table = apply_font(table, 10, 17)
    document.add_paragraph("")

    table = document.add_table(rows = 5, cols = 1)
    table.style = "Table Grid"
    cells = table.rows[0].cells
    run = cells[0].paragraphs[0].add_run("To Calculate Remaining Life in Years:")
    run.bold = True

    cells = table.rows[1].cells
    run = cells[0].paragraphs[0].add_run("Based on Current Structural Parameters")
    run.bold = True

    cells = table.rows[2].cells
    # cells[0].text = "Remaining Life (Years) = (Harmonic Mean Remaining ESALs) / (Typical Accumulated ESALs based on AADT)"
    run = cells[0].paragraphs[0].add_run("Remaining Life (Years) = ")
    run.bold = True
    run = cells[0].paragraphs[0].add_run("(Harmonic Mean Remaining ESALs)")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 173, 71)
    run = cells[0].paragraphs[0].add_run(" / ")
    run.bold = True
    run = cells[0].paragraphs[0].add_run("(Typical Accumulated ESALs based on AADT)")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)



    cells = table.rows[3].cells
    run = cells[0].paragraphs[0].add_run("Based on Current Rideability (IRI)")
    run.bold = True

    cells = table.rows[4].cells
    # cells[0].text = "Remaining Life (Years) = (Remaining ESALs, xx% Reliability) / (Typical Accumulated ESALs based on AADT)"
    run = cells[0].paragraphs[0].add_run("Remaining Life (Years) = ")
    run.bold = True
    run = cells[0].paragraphs[0].add_run("(Remaining ESALs, xx% Reliability)")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(0, 112, 192)
    run = cells[0].paragraphs[0].add_run(" / ")
    run.bold = True
    run = cells[0].paragraphs[0].add_run("(Typical Accumulated ESALs based on AADT)")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(112, 48, 160)

    table = apply_font(table, 10, 5)

    return document



def get_rp_str(rp_val, dmi_val):
    if(rp_val == None or dmi_val == None):
        return ""
    rp_str = str(rp_val) + " is FWD Station (DMI) " + str(dmi_val) +" Feet\n"
    return rp_str

def comments_page(ll, mde_path, comments, document):
    heading = "Comments for " + elab_dir(ll["dir"][:2]) + " from " + get_pathstring(mde_path)
    # document.add_heading(heading, 1)
    p = document.add_paragraph()
    sentence = p.add_run(heading)
    sentence.font.name = 'Times New Roman'
    sentence.font.size = docx.shared.Pt(14)
    sentence.font.underline = True
    sentence.bold = True
    sentence.font.color.rgb = docx.shared.RGBColor(112, 48, 160)

    document.add_paragraph(''.join(comments))
    return document

def gen_report(ll, mde, calc_data, stats_data, mde_path, f25_path, ll_no, year, con, special_case):
    comments_arr, dmi_val, rp_val = comments.get_comments(f25_path)
    report_page4.assign_values(rp_val, dmi_val)
    rp_str = get_rp_str(rp_val, dmi_val)
    document = report_page4.create_doc()
    document = report_page4.header_footer(document, mde_path, ll, "Underseal")
    document = cover_page(ll_no, year, mde_path, ll, mde, document, con) #page 1
    document.add_page_break()
    if(ll["pavtype"]!="asphalt"):
        document = underseal_page(rp_str, mde_path, ll, calc_data, document) #page 2
    document = chart(rp_str, stats_data, calc_data, mde, document)#page 3
    document = report_page4.soil_profile_page(document, mde, calc_data,rp_val,dmi_val)
    document.add_page_break()
    if(ll["pavtype"]!="asphalt"):
        document = report_page4.dm_page(document, calc_data, mde)
    document.add_page_break()
    document = esals_page(rp_str, calc_data, document)
    document = report_page4.overlay_design_page(document, ll, calc_data, report_page4.dir_str(ll["dir"][:2]), mde_path, post_design=False)

    document.add_page_break()

    #(to be modify)# change the "calc_data" in this line to "post underseal" data
    # Asphalt donesn't have post design page
    if(ll["pavtype"]!="asphalt" and (not special_case)):
        document = report_page4.overlay_design_page(document, ll, calc_data, report_page4.dir_str(ll["dir"][:2]), mde_path, post_design=True)
        document.add_page_break()

    document = comments_page(ll, mde_path, comments_arr, document)
    
    # Save the report in the same folder where MDE file is located
    report_save_folder = os.path.dirname(mde_path)
    report_filename = os.path.basename(mde_path)[:-4]
    # other_info = 222
    # report_filename ="FWD Results for {} {}".format(ll["req no"],other_info)
    docx_file_path = os.path.join(report_save_folder, "{}.docx".format(report_filename))
    document.save(docx_file_path)

    #########################################
    # # Convert to PDF session
    # # Save docuemnt as PDF and rename it
    # # Reopen the .docx file and save it as pdf
    # pdf_file_path = os.path.join(report_save_folder,"{}.pdf".format(report_filename))
    # convert(docx_file_path,pdf_file_path)
    #########################################

