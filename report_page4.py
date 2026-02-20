import docx as docx
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement, ns
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pickle as pkl
import numpy as np

from datetime import date

import re

rp_val = 0
dmi_val = 0

def assign_values(rp_val_tmp, dmi_val_tmp):
    global rp_val
    global dmi_val
    rp_val = rp_val_tmp
    dmi_val = dmi_val_tmp

######################################################
def m2ft(input_meter):
    """
    helper function to convert meter to feet
    """
    converted_feet = 3.28084*input_meter
    return converted_feet

def dir_str(dir):
    """Expand direction abbreviation to full text. Same as report.elab_dir()."""
    mapping = {"EB": "East Bound", "NB": "North Bound", "SB": "South Bound", "WB": "West Bound"}
    return mapping.get(dir, dir)

def get_from_rp_to_rp_str(mde_path, args):

    """
    This function will return 2 list. 
    The first list is the RP numbers of "from RP"
    The first list is the RP numbers of "to RP"
    For each list, the first element would be the refrence point number,
    the second element would be the milage number

    For example: 

    from_rp_list = ['142','-93']
    to_rp_list = ['142','62']

    """

    temp = os.path.splitext(os.path.basename(mde_path))[0]
    from_rp_match = re.findall(r'RP-(\d+)\+(-?\d+\.?\d?) to',temp)
    if len(from_rp_match) != 1 and not args.user_input:
        raise Exception('Something wrong when extracting the from RP!')
    try:
        from_rp_list = [from_rp_match[0][0],from_rp_match[0][1]]
    except (IndexError, TypeError):
        from_rp_list = ['Nan', 'Nan']

    to_rp_match = re.findall(r'to RP-(\d+)\+(-?\d+\.?\d?)',temp)
    if len(to_rp_match) != 1 and not args.user_input:
        raise Exception('Something wrong when extracting the to RP!')
    try:
        to_rp_list = [to_rp_match[0][0],to_rp_match[0][1]]
    except (IndexError, TypeError):
        to_rp_list = ['Nan', 'Nan']

    return from_rp_list, to_rp_list

# Header and footer related code
#################################################
def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(p):

    run = p.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText1 = create_element('w:instrText')
    create_attribute(instrText1, 'xml:space', 'preserve')
    instrText1.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText1)
    run._r.append(fldChar2)

    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def get_pathstring(mde_path, args):
    if args.user_input:
        return os.path.splitext(os.path.basename(mde_path))[0]
    else:
        tmp = os.path.splitext(os.path.basename(mde_path))[0]
        arr = tmp.split()
        arr[1] = "from"
        tmp = ' '.join(arr)
        return tmp


def header_footer(doc, mde_path, ll, args):
    today = date.today()
    d1 = today.strftime("%m/%d/%Y")

    paragraph = doc.sections[-1].header.paragraphs[0]
    p = paragraph._p  # ---this is the paragraph XML element---
    p.getparent().remove(p)

    table = doc.sections[-1].header.add_table(2,2,Inches(6))

    row_1_cell = table.rows[0].cells
    row_1_cell[0].paragraphs[0].add_run("R&D ID No: "+ll["req no"])
    row_1_cell[0].paragraphs[0].paragraph_format.line_spacing = 1
    row_1_cell[0].paragraphs[0].paragraph_format.space_after = Pt(0)
    add_page_number(row_1_cell[1].paragraphs[0])
    row_1_cell[1].paragraphs[0].paragraph_format.line_spacing = 1
    row_2_cell = table.rows[1].cells
    row_2_cell[0].paragraphs[0].add_run(d1)
    row_2_cell[1].paragraphs[0].paragraph_format.line_spacing = 1
    row_1_cell[1].paragraphs[0].paragraph_format.space_after = Pt(0)

    # Edit the footer section
    if(ll["pavtype"] == "asphalt"):
        treatment = "Overlay Design"
    elif(ll["pavtype"] == "concrete"):
        treatment = "Underseal"
    elif(ll["pavtype"] == "composite"):
        treatment = "Underseal and Overlay Design"
    section = doc.sections[-1]
    footer = section.footer
    paragraph = footer.add_paragraph()
    paragraph.text = get_pathstring(mde_path, args) + "\t\t" + treatment
    paragraph.style = doc.styles["footer"]
    # Edit the format of footer to avoid overflow of text due to filename is too long.
    sec = doc.sections[0]
    # finding end_point for the content 
    margin_end = docx.shared.Inches(
        sec.page_width.inches - (sec.left_margin.inches + sec.right_margin.inches))
    tab_stops = paragraph.paragraph_format.tab_stops
    # adding new tab stop, to the end point, and making sure that it's `RIGHT` aligned.
    tab_stops.add_tab_stop(margin_end, docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
    return doc
#################################################

def create_doc():
    document = docx.Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = docx.shared.Pt(12)

    return document

######################################################


def soil_profile_page(document,mde,calc_data,rp_val,dmi_val):

    document.add_paragraph(
        'Soil Profile around Joints and Cracks, {} is FWD Station (DMI) {} Feet'.format(rp_val,dmi_val)
        )

    table = document.add_table(rows=1,cols=6)
    # Set the style of the table
    table_cell_width = []
    table.style = 'Table Grid'
    # for col_idx,cell in enumerate(table.columns[0].cells):
    #     cell.width = Inches(table_cell_width[col_idx])

    hdr_cells = table.rows[0].cells
    run = hdr_cells[0].paragraphs[0].add_run('FWD Station (ft)')
    run.bold=True
    run = hdr_cells[1].paragraphs[0].add_run('Latitude')
    run.bold=True
    run = hdr_cells[2].paragraphs[0].add_run('Longitude')
    run.bold=True
    run = hdr_cells[3].paragraphs[0].add_run('Surface Deflection')
    run.bold=True
    run = hdr_cells[4].paragraphs[0].add_run('In-Situ CBR')
    run.bold=True
    run = hdr_cells[5].paragraphs[0].add_run('In-situ Structural Number')
    run.bold=True

    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Loop over all the rows
    for row_idx in range(calc_data['sensordata'].shape[0]):
        row_cells = table.add_row().cells
        # 3 here because we have 3 drops and we only take the second drop
        row_cells[0].text = str(round(m2ft(float(mde['deflections'][row_idx*3+1,1]))))
        # In case GPS is missing
        gps_1,gps_2 = mde['deflections'][row_idx*3+1,-6], mde['deflections'][row_idx*3+1,-5]
        if gps_1 and gps_2:
            row_cells[1].text = '{:.7f}'.format(float(gps_1)) # GPS
            row_cells[2].text = '{:.7f}'.format(float(gps_2)) # GPS
        else:
            row_cells[1].text = 'No GPS data' # GPS
            row_cells[2].text = 'No GPS data' # GPS
        row_cells[3].text = '{:.2f}'.format(calc_data['sensordata'][row_idx,0])
        row_cells[4].text = '{:.2f}'.format(calc_data['cbrarr'][row_idx])
        row_cells[5].text = '{:.2f}'.format(calc_data['sn'][row_idx])

        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    return document

def dm_page(document,calc_data,mde):

    document.add_paragraph(
        'Dynamic Modulus, {} is FWD Station (DMI) {} Feet'.format(rp_val,dmi_val)
        )

    table = document.add_table(rows=1,cols=2)
    # Set the style of the table
    table_cell_width = []
    table.style = 'Table Grid'
    # for col_idx,cell in enumerate(table.columns[0].cells):
    #     cell.width = Inches(table_cell_width[col_idx])

    hdr_cells = table.rows[0].cells
    run = hdr_cells[0].paragraphs[0].add_run('FWD Station (ft)')
    run.bold=True
    run = hdr_cells[1].paragraphs[0].add_run('Dynamic K-value of Concrete Pavement Support (pci)')
    run.bold=True

    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Loop over all the rows
    for row_idx in range(calc_data['sensordata'].shape[0]):
        row_cells = table.add_row().cells
        # 3 here because we have 3 drops and we only take the second drop
        row_cells[0].text = str(round(m2ft(float(mde['deflections'][row_idx*3+1,1]))))
        row_cells[1].text = str(round(calc_data['pcc_mod'][row_idx]))

        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    return document

def overlay_design_table(document, ll, calc_data, mde_path, dir_text, args, post_design=False):
    """
    Overlay design table and title
    """
    from_rp_list, to_rp_list = get_from_rp_to_rp_str(mde_path, args)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sentence = paragraph.add_run('{} Lane from RP {}+{} to RP {}+{}'.format(dir_text, from_rp_list[0], from_rp_list[1], to_rp_list[0], to_rp_list[1]))
    sentence.font.name = 'Arial'
    sentence.font.size = docx.shared.Pt(13)
    sentence.font.underline = True
    sentence.bold = True
    sentence.font.color.rgb = docx.shared.RGBColor(128, 0, 128)

    document.add_paragraph(
        'Overlay Requirement Profile, {} is FWD Station (DMI) {} Feet'.format(rp_val, dmi_val))

    table = document.add_table(rows=1,cols=3)
    # Set the style of the table
    table_cell_width = []
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    run = hdr_cells[0].paragraphs[0].add_run('In-Situ')
    run.bold=True
    run.font.color.rgb = docx.shared.RGBColor(26, 36, 238)
    run = hdr_cells[1].paragraphs[0].add_run('Average')
    hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.bold=True
    run.font.color.rgb = docx.shared.RGBColor(26, 36, 238)
    run = hdr_cells[2].paragraphs[0].add_run('Standard Deviation')
    hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.bold=True
    run.font.color.rgb = docx.shared.RGBColor(26, 36, 238)

    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def overlay_design_row(table,in_situ_str,average,sd):
        row_cells = table.add_row().cells
        run = row_cells[0].paragraphs[0].add_run(in_situ_str)
        run.bold=True
        row_cells[0].paragraphs[0].paragraph_format.line_spacing = 1.5
        row_cells[1].text = average # average
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].paragraphs[0].paragraph_format.line_spacing = 1.5
        row_cells[2].text = sd # standard
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].paragraphs[0].paragraph_format.line_spacing = 1.5

        return table

    def cal_avg_std(arr):
        """
        Always use the first column to calculate 'e_avg'????????????
        """
        arr_avg = np.mean(arr)
        arr_std = np.std(arr)

        # convert ksi to psi
        return arr_avg, arr_std

    if(post_design == False):
        e_avg, e_std = cal_avg_std(calc_data['e'][0,:])
        sn_avg, sn_std = cal_avg_std(calc_data['sn'])
        mr_avg, mr_std = cal_avg_std(calc_data['insitumr'])
        cbr_avg, cbr_std = cal_avg_std(calc_data['cbrarr'])
        if(ll["pavtype"]!="asphalt"):
            dyn_k_avg, dyn_k_std = cal_avg_std(np.array(calc_data['pcc_mod']))
        if(ll["pavtype"]=="composite"):
            e2_avg, e2_std = cal_avg_std(calc_data['e'][1,:]) # e2_avg would be concrete in the case of composite
    else:
        if(ll["pavtype"]=="composite"):
            e_avg, e_std = cal_avg_std(calc_data["new_stats"]['e1'])
            e2_avg, e2_std = cal_avg_std(calc_data["new_stats"]['e2'])
        elif (ll["pavtype"]=="concrete"):
            e_avg, e_std = cal_avg_std(calc_data["new_stats"]['e2'])
        else:
            raise Exception("Error in the source code, pavement type other than 'composite' and 'asphalt' should not have post design page.")
        sn_avg, sn_std = cal_avg_std(calc_data["new_stats"]['sn'])
        mr_avg, mr_std = cal_avg_std(calc_data["new_stats"]['mr'])
        cbr_avg, cbr_std = cal_avg_std(calc_data["new_stats"]['cbr'])
        if(ll["pavtype"]!="asphalt"):
            dyn_k_avg, dyn_k_std = cal_avg_std(np.array(calc_data["new_stats"]['k']))
    # rxn_sungrade is list??????????????????

    if ll["pavtype"]=="concrete":
        overlay_design_row(table,
                           'Elastic Modulus of Concrete (psi)', 
                           '{:,}'.format(round(e_avg*1e3)), 
                           '{:,}'.format(round(e_std*1e3)))
    elif ll["pavtype"]=="asphalt":
        overlay_design_row(table,
                           'Elastic Modulus of Asphalt (psi)', 
                           '{:,}'.format(round(e_avg*1e3)), 
                           '{:,}'.format(round(e_std*1e3)))
    elif ll["pavtype"]=="composite":
        overlay_design_row(table,
                           'Elastic Modulus of Asphalt (psi)', 
                           '{:,}'.format(round(e_avg*1e3)), 
                           '{:,}'.format(round(e_std*1e3)))
        overlay_design_row(table,
                           'Elastic Modulus of Concrete (psi)', 
                           '{:,}'.format(round(e2_avg*1e3)), 
                           '{:,}'.format(round(e2_std*1e3)))
    
    overlay_design_row(table,
                       'Structural Number',
                       '{:.2f}'.format(sn_avg), 
                       '{:.2f}'.format(sn_std))
    overlay_design_row(table,
                       'Modulus of Resilient of Subgrade Soil (psi)',
                       '{:,}'.format(round(mr_avg)), 
                       '{:,}'.format(round(mr_std)))
    overlay_design_row(table, 
                       'CBR of Subgrade Soil (%)', 
                       '{:.2f}'.format(cbr_avg), 
                       '{:.2f}'.format(cbr_std))
    if(ll["pavtype"]!="asphalt"):
        overlay_design_row(table,
                           'Dynamic K-value of Pavement Support (pci)',
                           '{:,}'.format(round(dyn_k_avg)), 
                           '{:,}'.format(round(dyn_k_std)))

    return document

def overlay_design_page(document, ll, calc_data, dir_text, mde_path, args, post_design=False):

    paragraph = document.add_paragraph()
    if post_design:
        sentence = paragraph.add_run('Estimated Overlay Design after Underseal')
    else:
        sentence = paragraph.add_run('Overlay Design')
    
    sentence.font.name = 'Arial'
    sentence.font.size = docx.shared.Pt(14)
    sentence.bold = True
    sentence.italic = True
    sentence.font.color.rgb = docx.shared.RGBColor(26, 36, 238)

    document = overlay_design_table(document, ll, calc_data, mde_path, dir_text, args, post_design)

    return document